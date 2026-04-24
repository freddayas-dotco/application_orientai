"""
OrientAI — Rapport personnalisé v2
Améliorations vs v1 :
- Page de couverture soignée
- Table des matières
- Introduction personnalisée
- Graphique radar des soft skills (matplotlib → image)
- Top 15 métiers recommandés
- Cases à cocher vides (☐) pour que l'élève remplisse lui-même
"""

import io
import math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Données partagées importées depuis donnees_rapport.py ─────────────────────
from donnees_rapport import MAPPING_QUESTIONS, SS_RAPPORT_NOMS, SS_NIVEAUX, SS_COULEURS_HEX, SS_COULEURS_RGB


INTRO_NIVEAU = {
    "fort": "Votre profil révèle des soft skills globalement très développées. Vous disposez d'atouts solides pour intégrer des métiers exigeants et évoluer rapidement.",
    "moyen": "Votre profil montre un socle de soft skills équilibré, avec des points forts bien identifiés et des axes de progression concrets à exploiter.",
    "faible": "Votre profil montre des marges de progression importantes. C'est une opportunité : les soft skills se travaillent et s'améliorent avec de la pratique et de la motivation.",
}


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _para(doc, text="", bold=False, size=11, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def generer_radar_chart(ss_vals):
    """Génère un graphique radar des 8 soft skills → retourne bytes PNG."""
    labels = [
        "Communication", "Esprit\ncritique", "Éthique",
        "Intel.\némotionnelle", "Intel.\nsociale",
        "Mgmt\nprojet", "Mgmt\néquipe", "Organisation"
    ]
    N = len(labels)
    angles = [n / float(N) * 2 * math.pi for n in range(N)]
    angles += angles[:1]

    values = [v / 3 for v in ss_vals]
    values += values[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.set_facecolor('#fafaf8')
    fig.patch.set_facecolor('#fafaf8')

    # Zones de fond par niveau
    for level, alpha, color in [(1/3, 0.15, '#e8a87c'), (2/3, 0.1, '#c9a84c'), (1.0, 0.07, '#6b9e7e')]:
        ax.fill([a for a in angles], [level] * (N+1), alpha=alpha, color=color)

    # Contour principal
    ax.plot(angles, values, 'o-', linewidth=2, color='#1a1a2e', markersize=5)
    ax.fill(angles, values, alpha=0.25, color='#1a1a2e')

    # Configuration axes
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, size=8, color='#1a1a2e')
    ax.set_ylim(0, 1)
    ax.set_yticks([1/3, 2/3, 1])
    ax.set_yticklabels(["À développer", "Intermédiaire", "Très développée"], size=7, color='#666')
    ax.grid(color='#cccccc', linestyle='--', linewidth=0.5)
    ax.spines['polar'].set_color('#cccccc')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#fafaf8')
    plt.close()
    buf.seek(0)
    return buf


def generer_rapport_v2(prenom: str, nom: str, answers: list, profile: dict, top_metiers: list = None) -> bytes:
    doc = Document()

    # ── Marges ────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    ss_vals = profile["ss"]
    score_moyen = sum(ss_vals) / len(ss_vals)
    niveau_global = "fort" if score_moyen >= 2.5 else "moyen" if score_moyen >= 1.8 else "faible"

    points_forts = [SS_RAPPORT_NOMS[i] for i, v in enumerate(ss_vals) if v == 3]
    a_developper = [SS_RAPPORT_NOMS[i] for i, v in enumerate(ss_vals) if v == 1]

    # ══════════════════════════════════════════════════════════════════════════
    # PAGE DE COUVERTURE
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    # Logo/titre
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("🧭 OrientAI")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    sous = doc.add_paragraph()
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sous.add_run("Schéma de personnalité")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x7a, 0x7a, 0x8c)

    for _ in range(3):
        doc.add_paragraph()

    # Nom de l'élève
    nom_para = doc.add_paragraph()
    nom_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = nom_para.add_run(f"{prenom} {nom}")
    r3.font.size = Pt(26)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0xd4, 0x66, 0x7a)

    for _ in range(2):
        doc.add_paragraph()

    # Résumé rapide en bas de couverture
    if points_forts:
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = pf.add_run(f"Points forts : {', '.join(points_forts)}")
        r4.font.size = Pt(11)
        r4.font.color.rgb = RGBColor(0x6b, 0x9e, 0x7e)

    if a_developper:
        ad = doc.add_paragraph()
        ad.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r5 = ad.add_run(f"À développer : {', '.join(a_developper)}")
        r5.font.size = Pt(11)
        r5.font.color.rgb = RGBColor(0xe8, 0xa8, 0x7c)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE DES MATIÈRES (titre — Word la génère avec Ctrl+A puis F9)
    # ══════════════════════════════════════════════════════════════════════════
    toc_titre = doc.add_paragraph("Table des matières", style='Heading 1')
    toc_titre.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Champ TOC Word
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run4._r.append(fldChar3)

    toc_note = doc.add_paragraph()
    r_note = toc_note.add_run("→ Pour afficher la table des matières : Ctrl+A puis F9 dans Word")
    r_note.font.size = Pt(9)
    r_note.font.italic = True
    r_note.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # I. INTRODUCTION PERSONNALISÉE
    # ══════════════════════════════════════════════════════════════════════════
    h1 = doc.add_paragraph("I. Votre profil en un coup d'œil", style='Heading 1')
    h1.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    intro_p = doc.add_paragraph()
    intro_p.add_run(f"{prenom}, ").font.bold = True
    intro_text = INTRO_NIVEAU[niveau_global]
    intro_p.add_run(intro_text)
    intro_p.paragraph_format.space_after = Pt(10)

    # Résumé tableau soft skills
    h2 = doc.add_paragraph("Résumé de vos soft skills", style='Heading 2')
    h2.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x55)

    table_res = doc.add_table(rows=1, cols=2)
    table_res.style = 'Table Grid'
    hdr = table_res.rows[0].cells
    hdr[0].text = "Soft skill"
    hdr[1].text = "Niveau"
    for cell in hdr:
        _set_cell_bg(cell, "1A1A2E")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(10)

    for i, nom_ss in SS_RAPPORT_NOMS.items():
        niveau = ss_vals[i]
        row = table_res.add_row().cells
        row[0].text = nom_ss
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].text = SS_NIVEAUX[niveau]
        _set_cell_bg(row[1], SS_COULEURS_HEX[niveau])
        for p in row[1].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(10)

    for row in table_res.rows:
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(6)

    doc.add_paragraph()

    # ── Graphique radar ────────────────────────────────────────────────────────
    h2_radar = doc.add_paragraph("Diagramme de personnalité", style='Heading 2')
    h2_radar.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x55)

    radar_buf = generer_radar_chart(ss_vals)
    doc.add_picture(radar_buf, width=Cm(12))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    legende = doc.add_paragraph()
    legende.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, color in [("À développer", (0xe8, 0xa8, 0x7c)),
                          ("  Intermédiaire", (0xc9, 0xa8, 0x4c)),
                          ("  Très développée", (0x6b, 0x9e, 0x7e))]:
        r = legende.add_run(f"● {label}")
        r.font.color.rgb = RGBColor(*color)
        r.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # II. TOP 15 MÉTIERS
    # ══════════════════════════════════════════════════════════════════════════
    if top_metiers:
        h1_met = doc.add_paragraph("II. Vos métiers recommandés", style='Heading 1')
        h1_met.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        intro_met = doc.add_paragraph(
            f"Voici les 15 métiers les plus compatibles avec votre profil de soft skills, "
            f"classés par score de compatibilité."
        )
        intro_met.paragraph_format.space_after = Pt(10)

        table_met = doc.add_table(rows=1, cols=4)
        table_met.style = 'Table Grid'
        hdrs = table_met.rows[0].cells
        for txt, cell in zip(["Rang", "Métier", "Secteur", "Score"], hdrs):
            cell.text = txt
            _set_cell_bg(cell, "1A1A2E")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.bold = True
                    r.font.size = Pt(10)

        for rank, m in enumerate(top_metiers[:15], 1):
            score = m.get("score", 0)
            couleur = "6B9E7E" if score >= 75 else "C9A84C" if score >= 50 else "E8A87C"
            row = table_met.add_row().cells
            row[0].text = str(rank)
            row[1].text = m.get("metier", "")
            row[2].text = m.get("secteur", "").split(" ; ")[0]
            row[3].text = f"{score}%"
            _set_cell_bg(row[3], couleur)
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
            for p in row[3].paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.bold = True

        for row in table_met.rows:
            row.cells[0].width = Cm(1.5)
            row.cells[1].width = Cm(7)
            row.cells[2].width = Cm(5)
            row.cells[3].width = Cm(2)

        mention = doc.add_paragraph(
            "💡 Salaires indicatifs — Sources : ONISEP, APEC, INSEE, observatoires sectoriels. "
            "Les fourchettes peuvent varier selon l'expérience, la région et la taille de l'entreprise."
        )
        mention.runs[0].font.size = Pt(8)
        mention.runs[0].font.italic = True
        mention.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        mention.paragraph_format.space_before = Pt(6)

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # III. ANALYSE DÉTAILLÉE PAR SOFT SKILL
    # ══════════════════════════════════════════════════════════════════════════
    section_num = "III" if top_metiers else "II"
    h1_det = doc.add_paragraph(f"{section_num}. Analyse détaillée de vos soft skills", style='Heading 1')
    h1_det.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Mapping questions → SS
    ss_questions = {i: [] for i in range(8)}
    for q_idx, mapping in enumerate(MAPPING_QUESTIONS):
        sous_comp, soft_skill_nom, textes = mapping
        ss_idx = next(
            (i for i, nom_ss in SS_RAPPORT_NOMS.items()
             if nom_ss.lower() in soft_skill_nom.lower() or soft_skill_nom.lower() in nom_ss.lower()),
            None
        )
        if ss_idx is not None:
            ss_questions[ss_idx].append((q_idx, sous_comp, textes))

    lettres = ["A", "B", "C", "D"]

    for ss_idx, ss_nom in SS_RAPPORT_NOMS.items():
        niveau = ss_vals[ss_idx]
        couleur_hex = SS_COULEURS_HEX[niveau]
        couleur_rgb = SS_COULEURS_RGB[niveau]
        niveau_texte = SS_NIVEAUX[niveau]

        # Titre section
        h2_ss = doc.add_paragraph(f"Critère « {ss_nom} » : {niveau_texte}", style='Heading 2')
        h2_ss.runs[0].font.color.rgb = RGBColor(*couleur_rgb)

        doc.add_paragraph()

        # Tableau sous-compétences
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for txt, cell in zip(["Sous-compétence", "Ce que vos réponses révèlent", "Vous confirmez ?"], hdr_cells):
            cell.text = txt
            _set_cell_bg(cell, "2C2C2A")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.bold = True
                    r.font.size = Pt(10)

        qs = ss_questions.get(ss_idx, [])
        for row_idx, (q_idx, sous_comp, textes) in enumerate(qs):
            ans_idx = answers[q_idx] if q_idx < len(answers) and answers[q_idx] is not None else 0
            lettre = lettres[ans_idx]
            texte = textes.get(lettre, "")

            row = table.add_row().cells

            row[0].text = sous_comp
            for p in row[0].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.bold = True

            row[1].text = texte
            for p in row[1].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

            # Case à cocher VIDE pour que l'élève réponde
            row[2].text = "☐ Oui     ☐ Non"
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for p in row[2].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

            if row_idx % 2 == 0:
                _set_cell_bg(row[0], "F5F5F5")
                _set_cell_bg(row[1], "F5F5F5")
                _set_cell_bg(row[2], "F5F5F5")

        for row in table.rows:
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(10)
            row.cells[2].width = Cm(3)

        doc.add_paragraph()
        if ss_idx < 7:
            doc.add_page_break()

    # ── Pied de page ──────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = footer_p.add_run("Rapport généré par OrientAI v5.1 — Outil d'orientation par soft skills")
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
